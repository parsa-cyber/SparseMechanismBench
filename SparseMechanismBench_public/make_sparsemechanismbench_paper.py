from pathlib import Path
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT=Path(__file__).resolve().parent
RES=ROOT/'results'
FIG=ROOT/'figures'
OUT=ROOT/'SparseMechanismBench_manuscript.docx'


def pct(x):
    return f"{100*x:.1f}%"

def mean_pm(m,s):
    return f"{100*m:.1f}% ± {100*s:.1f}%"

def style_doc(doc):
    sec=doc.sections[0]
    sec.top_margin=Inches(0.65); sec.bottom_margin=Inches(0.65); sec.left_margin=Inches(0.7); sec.right_margin=Inches(0.7)
    styles=doc.styles
    styles['Normal'].font.name='Arial'; styles['Normal'].font.size=Pt(10)
    for s in ['Heading 1','Heading 2','Heading 3']:
        styles[s].font.name='Arial'; styles[s].font.bold=True
    styles['Heading 1'].font.size=Pt(16)
    styles['Heading 2'].font.size=Pt(13)
    styles['Heading 3'].font.size=Pt(11)

def add_caption(doc, text):
    p=doc.add_paragraph(text)
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.italic=True; r.font.size=Pt(9)

def set_cell_shading(cell, fill='D9EAF7'):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def add_table(doc, headers, rows, caption=None, widths=None):
    table=doc.add_table(rows=1, cols=len(headers))
    table.style='Table Grid'; table.alignment=WD_TABLE_ALIGNMENT.CENTER
    hdr=table.rows[0].cells
    for i,h in enumerate(headers):
        hdr[i].text=str(h)
        set_cell_shading(hdr[i])
        for p in hdr[i].paragraphs:
            for r in p.runs: r.bold=True; r.font.size=Pt(8.5)
    for row in rows:
        cells=table.add_row().cells
        for i,val in enumerate(row):
            cells[i].text=str(val)
            cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP
            for p in cells[i].paragraphs:
                for r in p.runs: r.font.size=Pt(8.5)
    if caption: add_caption(doc, caption)
    return table

def add_eq(doc, text):
    p=doc.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(text)
    r.font.name='Cambria Math'; r.font.size=Pt(12)

# Load CSVs
classification=pd.read_csv(RES/'updated_classification_summary.csv')
stdcnn=pd.read_csv(RES/'standard_cnn_summary.csv')
homeo=pd.read_csv(RES/'homeostatic_sparse_oja_summary.csv')
sweep=pd.read_csv(RES/'sparsity_sweep_subset_summary.csv')
replay=pd.read_csv(RES/'replay_sweep_subset_summary.csv')
stat=pd.read_csv(RES/'statistical_tests_summary.csv')
readout=pd.read_csv(RES/'frozen_readout_summary.csv')
geo=pd.read_csv(RES/'representation_geometry_summary.csv')

# Build doc
doc=Document(); style_doc(doc)

title='SparseMechanismBench: Why Sparse Local Plasticity Is Not Enough for Task-Discriminative Learning'
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run(title); r.bold=True; r.font.size=Pt(18)
p=doc.add_paragraph('Parsa Fatehi / Langley High School'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER

# Abstract
doc.add_heading('Abstract', level=1)
doc.add_paragraph(
"Modern artificial neural networks usually learn through global gradient-based optimization, while biological nervous systems rely heavily on local synaptic plasticity, sparse activity, competition, homeostasis, replay, and feedback. This project introduces SparseMechanismBench, a mechanism-isolation framework that asks whether sparse local-plasticity representations are weak because the representations themselves are poorly structured or because their supervised readouts are too limited. Across completed digits, MNIST, and Fashion-MNIST benchmarks, supervised/backpropagation models and a LeNet-style Standard CNN substantially outperform simple Oja-style feature learners. Sparse Oja reliably creates high activation sparsity, but does not produce competitive task-discriminative representations. A Homeostatic Sparse Oja variant improves some local-learning trade-offs but does not close the gap to backpropagation. Sparsity and replay sweeps show that extreme sparsity weakens task discrimination and that replay can reduce forgetting only partially. A new Frozen Feature Readout Experiment and Representation Geometry Analysis were implemented and completed on the digits dataset; MNIST/Fashion-MNIST reruns are blocked in the submitted package because raw IDX data are not included. The central conclusion is that sparsity is useful but insufficient: efficient biological learning likely requires coordinated interaction among local plasticity, sparse coding, competition, homeostasis, replay, feedback, and task-relevant signals.")

# Intro
doc.add_heading('1. Introduction', level=1)
doc.add_paragraph(
"A central difference between biological and artificial learning is the credit-assignment mechanism. Backpropagation updates parameters using gradients of a task loss, while local biological plasticity rules update synapses using variables available at or near the synapse. Sparse firing is often proposed as a bridge between biological efficiency and artificial learning, because sparse activity can reduce redundancy, metabolic cost, and interference. However, sparse activity is not the same as a task-discriminative representation. A representation can be sparse while failing to separate classes, support linear readout, or retain prior tasks.")
doc.add_paragraph(
"The earlier version of this project compared supervised baselines, MLPs, compact CNNs, Oja, and Sparse Oja on digits, MNIST, and Fashion-MNIST. It found a clean negative result: backpropagation-trained/supervised models outperformed simple Hebbian/Oja feature learners, while Sparse Oja produced high sparsity but poor task discrimination. The present revision turns that benchmark into a mechanism-isolation study. Instead of asking whether local learning is simply better or worse than backpropagation, SparseMechanismBench asks which biological mechanisms are sufficient or insufficient for closing the gap: local plasticity, sparse coding, competition, homeostasis, replay, and readout strength.")

doc.add_heading('Scientific Contribution', level=2)
doc.add_paragraph(
"This study makes four contributions. First, it separates representational sparsity from task-discriminative accuracy by measuring both across local and gradient-based learning systems. Second, it evaluates whether sparse local plasticity improves sample efficiency and continual-learning stability across digits, MNIST, and Fashion-MNIST. Third, it tests whether biologically inspired stabilizing mechanisms such as competition, adaptive homeostasis, and replay improve the accuracy-sparsity-forgetting trade-off. Fourth, it adds readout and representation-geometry diagnostics to test whether local features fail because the features are intrinsically poorly structured or merely hard to decode with a linear classifier. The main result is a strengthened negative result: sparse local plasticity creates sparse representations, but by itself it does not reliably create robust task-discriminative or stable representations.")

# Related Work
doc.add_heading('2. Related Work', level=1)
doc.add_heading('A. Biological local learning', level=2)
doc.add_paragraph(
"Hebbian learning formalizes the intuition that correlated pre- and postsynaptic activity should strengthen a connection (Hebb, 1949). Oja's rule adds a stabilizing normalization term that prevents unbounded weight growth and can recover principal components in a simplified neuron model (Oja, 1982). Spike-timing-dependent plasticity extends this local idea to temporal order, with synaptic changes depending on the relative timing of pre- and postsynaptic spikes (Bi & Poo, 1998). BCM theory adds a sliding activity-dependent threshold for potentiation versus depression, making selectivity depend on a neuron's recent activity history (Bienenstock, Cooper, & Munro, 1982). Predictive coding models go beyond purely local correlation by introducing top-down predictions and error signals across cortical hierarchies (Rao & Ballard, 1999). These theories are biologically plausible because they use local activity, timing, or recurrent error signals, but simple local plasticity alone lacks the direct task-directed credit assignment that supervised learning provides.")
doc.add_heading('B. Sparse coding and efficient representations', level=2)
doc.add_paragraph(
"Sparse coding proposes that sensory systems represent inputs using only a small subset of active units. Olshausen and Field (1996, 1997) showed that sparse coding of natural images can learn localized, oriented filters resembling V1 simple-cell receptive fields. Sparse autoencoders and k-winners-take-all mechanisms similarly encourage only a limited number of hidden units to activate for each input (Makhzani & Frey, 2013). Competitive learning and lateral inhibition provide biologically motivated routes to sparse distributed codes. However, sparsity is not equivalent to class separability. A feature space can be metabolically efficient and still poorly aligned with a downstream classification objective. This distinction motivates the core hypothesis of SparseMechanismBench: sparsity may be necessary for biological efficiency but insufficient for task-discriminative learning.")
doc.add_heading('C. Biologically plausible credit assignment', level=2)
doc.add_paragraph(
"Backpropagation remains the dominant algorithm for training artificial neural networks (Rumelhart, Hinton, & Williams, 1986), but it is often criticized as biologically implausible because it requires precise gradients, nonlocal error propagation, and weight-transport-like assumptions. Feedback alignment shows that exact symmetric feedback weights are not always necessary for useful credit assignment (Lillicrap et al., 2016). Equilibrium propagation learns in recurrent energy-based systems through a free phase and a weakly clamped phase, offering a more biologically plausible approximation to gradient-based learning (Scellier & Bengio, 2017). Target propagation uses layerwise targets rather than gradients (Lee et al., 2015). Predictive coding and dendritic-error models propose additional routes through which local circuits might approximate global credit assignment (Whittington & Bogacz, 2019). These approaches suggest that the weakness of simple Oja/Sparse Oja may not be local learning per se, but the absence of task-relevant feedback.")
doc.add_heading('D. Catastrophic forgetting and replay', level=2)
doc.add_paragraph(
"Continual learning highlights another limitation of purely gradient-based models: new-task training can overwrite old-task representations. Early connectionist work described catastrophic interference in sequential learning (McCloskey & Cohen, 1989). Complementary learning systems theory argues that the hippocampus and neocortex support different learning timescales, with replay helping stabilize long-term knowledge (McClelland, McNaughton, & O'Reilly, 1995). Experience replay has also become central in reinforcement learning and deep reinforcement learning (Lin, 1992; Mnih et al., 2015). Elastic weight consolidation reduces forgetting by penalizing changes to parameters important for previous tasks (Kirkpatrick et al., 2017). Replay is biologically and computationally motivated, but its interpretation depends on initial learning: a model that never learns Task 1 cannot meaningfully demonstrate strong retention.")
doc.add_heading('E. Why this work is different', level=2)
doc.add_paragraph(
"Most prior studies focus on one mechanism: a local rule, a sparsity objective, a credit-assignment approximation, or a replay strategy. SparseMechanismBench instead places these mechanisms in one reproducible pipeline and separates the quantities that are often conflated: sparsity, sample efficiency, task-discriminative accuracy, representation geometry, readout strength, computational cost, and continual-learning retention. The goal is not to claim that one toy local rule explains brain-like learning. The goal is to identify which mechanisms are insufficient alone and which combinations move local learning closer to useful task-directed representation.")

# Research question
doc.add_heading('3. Research Question and Hypotheses', level=1)
doc.add_paragraph("Central question: Are sparse local-plasticity representations weak because the representations themselves are poor, or because the downstream readout/classifier is too weak?")
for h in [
"H1: Supervised/backpropagation-based models will outperform simple Hebbian/Oja feature learners in accuracy and sample efficiency.",
"H2: Sparse Oja-style models will produce substantially higher activation sparsity than dense MLPs.",
"H3: Competition and homeostasis may improve the accuracy-sparsity trade-off relative to fixed Sparse Oja, but will not necessarily close the supervised learning gap.",
"H4: Stronger frozen readouts can diagnose whether local features contain task information that is not linearly accessible.",
"H5: Replay may reduce raw forgetting, but retention is meaningful only when Task 1 is first learned well." ]:
    doc.add_paragraph(h, style='List Bullet')

# Math setup
doc.add_heading('4. Mathematical Setup', level=1)
doc.add_paragraph("For presynaptic activity x_j, postsynaptic activity y_i, and learning rate eta, a Hebbian update is:")
add_eq(doc, "Δwᵢⱼ = η xⱼ yᵢ")
doc.add_paragraph("Oja's local normalization modifies this update as:")
add_eq(doc, "Δwᵢ = η(yᵢx − yᵢ²wᵢ)")
doc.add_paragraph("Backpropagation updates a parameter w by descending the gradient of a task loss L:")
add_eq(doc, "w ← w − η ∂L/∂w")
doc.add_paragraph("Raw and normalized forgetting are defined as:")
add_eq(doc, "F = A_before − A_after")
add_eq(doc, "F_norm = (A_before − A_after) / A_before")
doc.add_paragraph("A_before is Task 1 accuracy before Task 2 training, and A_after is Task 1 accuracy after Task 2 training. Normalized forgetting is interpreted cautiously when A_before is near chance.")

# Methods
doc.add_heading('5. Methods', level=1)
doc.add_heading('Datasets', level=2)
# dataset table from dataset_validation
try:
    dv=pd.read_csv(RES/'dataset_validation.csv')
    rows=[]
    for _,r in dv.iterrows(): rows.append([r['dataset'], r['train_size'], r['test_size'], r['image_shape'], f"{r['train_min']:.1f}-{r['train_max']:.1f}", f"{r['test_min']:.1f}-{r['test_max']:.1f}"])
    add_table(doc, ['Dataset','Train','Test','Shape','Train range','Test range'], rows, 'Table 1. Dataset validation. Source: dataset_validation.csv.')
except Exception:
    doc.add_paragraph('Dataset validation table unavailable.')
if (FIG/'example_images_grid.png').exists():
    doc.add_picture(str(FIG/'example_images_grid.png'), width=Inches(5.5)); add_caption(doc, 'Figure 1. Example images from digits, MNIST, and Fashion-MNIST. Source: validate_datasets.py.')

doc.add_heading('Models and Mechanisms', level=2)
add_table(doc, ['Model','Local?','Sparse?','Homeostatic?','Replay?','Task-directed signal'], [
['Oja','Yes','No','No','Optional','No'],['Sparse Oja','Yes','Yes','No','Optional','No'],['Homeostatic Sparse Oja','Yes','Yes','Yes','Optional','No'],['MLP','No','No','No','Optional','Yes'],['Standard CNN','No','No','No','Optional','Yes'],['Frozen readouts','No','Depends on features','No','No','Yes readout only']], 'Table 2. Mechanism ladder tested in SparseMechanismBench.')

doc.add_heading('Frozen Feature Readout Experiment', level=2)
doc.add_paragraph(
"Oja, Sparse Oja, and Homeostatic Sparse Oja were trained without labels on the training split. Their features were then frozen, and several supervised readouts were trained on top: logistic regression, ridge classifier, linear SVM, k-nearest neighbors, and a small MLP readout. Raw pixels, supervised MLP hidden features, and Standard CNN penultimate features served as comparison representations. In the current package this experiment completed on digits. MNIST/Fashion-MNIST execution is implemented but blocked because the submitted reproducibility ZIP does not include raw IDX files.")

doc.add_heading('Representation Geometry Analysis', level=2)
doc.add_paragraph(
"For each representation, the analysis computes linear probe accuracy, silhouette score, Davies-Bouldin index, Fisher between/within scatter ratio, mean intra-class distance, mean inter-class distance, inter/intra distance ratio, activation sparsity, dead-unit rate, mean activation, utilization entropy, and class selectivity. PCA visualizations are generated with a fixed seed. In the current package this analysis completed on digits and is scripted for MNIST/Fashion-MNIST once raw IDX files are provided.")

# Results
doc.add_heading('6. Results', level=1)
doc.add_heading('A. Supervised baselines vs. local plasticity', level=2)
models=['logistic_regression','mlp','mlp_dropout','compact_cnn','standard_cnn','hebbian_oja','sparse_hebbian_oja','homeostatic_sparse_oja']
model_labels={'logistic_regression':'Logistic regression','mlp':'MLP','mlp_dropout':'MLP + dropout','compact_cnn':'compact CNN','standard_cnn':'Standard CNN','hebbian_oja':'Oja','sparse_hebbian_oja':'Sparse Oja','homeostatic_sparse_oja':'Homeostatic Sparse Oja'}
rows=[]
for m in models:
    row=[model_labels[m]]
    for ds in ['digits','mnist','fashion_mnist']:
        g=classification[(classification.dataset==ds)&(classification.model==m)]
        if len(g)==0:
            if m=='standard_cnn':
                g=stdcnn[(stdcnn.dataset==ds)&(stdcnn.model=='standard_cnn')]
                if len(g): row.append(mean_pm(g.iloc[0]['accuracy_mean'],g.iloc[0]['accuracy_std']))
                else: row.append('not run')
            elif m=='homeostatic_sparse_oja':
                g=homeo[(homeo.dataset==ds)&(homeo.model=='homeostatic_sparse_oja')]
                if len(g): row.append(mean_pm(g.iloc[0]['accuracy_mean'],g.iloc[0]['accuracy_std']))
                else: row.append('not run')
            else: row.append('not run')
        else: row.append(mean_pm(g.iloc[0]['accuracy_mean'],g.iloc[0]['accuracy_std']))
    rows.append(row)
add_table(doc,['Model','digits','MNIST','Fashion-MNIST'],rows,'Table 3. Updated classification accuracy. Source: updated_classification_summary.csv, standard_cnn_summary.csv, and homeostatic_sparse_oja_summary.csv.')
for img,cap in [('standard_vs_compact_cnn.png','Figure 2. Standard CNN resolves the earlier weak-CNN concern. Source: standard_cnn_by_seed.csv and updated_classification_summary.csv.'),('mechanism_ladder_accuracy.png','Figure 3. Mechanism ladder classification results. Source: updated_classification_summary.csv.')]:
    if (FIG/img).exists(): doc.add_picture(str(FIG/img), width=Inches(5.8)); add_caption(doc, cap)

doc.add_heading('B. Sparsity-accuracy trade-off', level=2)
rows=[]
for _,r in sweep.iterrows():
    if r['dataset'] in ['mnist','fashion_mnist'] and r['model'] in ['Sparse Oja','Oja']:
        rows.append([r['dataset'].replace('_','-'), f"{100*r['active_frac']:.0f}%" if pd.notna(r['active_frac']) else '100%', mean_pm(r['sparsity_mean'],r['sparsity_std']), mean_pm(r['accuracy_mean'],r['accuracy_std'])])
add_table(doc,['Dataset','Active units','Sparsity','Accuracy'],rows,'Table 4. Sparsity sweep on fixed subset. Source: sparsity_sweep_subset_summary.csv.')
if (FIG/'ablation_accuracy_vs_sparsity.png').exists(): doc.add_picture(str(FIG/'ablation_accuracy_vs_sparsity.png'), width=Inches(6.2)); add_caption(doc,'Figure 4. Sparsity sweep: extreme sparsity weakens task discrimination; accuracy improves as more units are active. Source: sparsity_sweep_subset_summary.csv.')

doc.add_heading('C. Frozen Feature Readout Experiment', level=2)
best = readout[readout['dataset']=='digits'].loc[readout[readout['dataset']=='digits'].groupby('feature_source')['accuracy_mean'].idxmax()].copy()
order=['raw_pixels','oja','sparse_oja','homeostatic_sparse_oja','mlp_hidden','standard_cnn_penultimate']
best['feature_source']=pd.Categorical(best['feature_source'],categories=order,ordered=True); best=best.sort_values('feature_source')
rows=[]
for _,r in best.iterrows():
    rows.append([r['feature_source'].replace('_',' '), r['readout'], int(r['n']), mean_pm(r['accuracy_mean'],r['accuracy_std']), mean_pm(r['feature_sparsity_mean'],r['feature_sparsity_std'])])
add_table(doc,['Frozen feature source','Best readout','Seeds','Accuracy','Feature sparsity'],rows,'Table 5. Frozen readout experiment on digits. Source: frozen_readout_summary.csv.')
doc.add_paragraph(
"Digits results show that readout strength matters: Oja features are highly decodable with a linear SVM (97.5% ± 0.4%), close to raw-pixel KNN (98.3%). Sparse Oja and Homeostatic Sparse Oja remain much weaker at roughly 83% despite stronger readouts, suggesting that extreme sparsity can damage accessible class information. MNIST/Fashion-MNIST frozen-readout results are not reported because the raw IDX data are absent from the submitted ZIP; the code records this as a blocked experiment rather than fabricating values.")
if (FIG/'frozen_readout_digits_heatmap.png').exists(): doc.add_picture(str(FIG/'frozen_readout_digits_heatmap.png'), width=Inches(6.2)); add_caption(doc,'Figure 5. Frozen readout accuracy on digits. Source: frozen_readout_summary.csv.')

doc.add_heading('D. Representation Geometry', level=2)
rows=[]
for _,r in geo[geo.dataset=='digits'].iterrows():
    rows.append([r['source'].replace('_',' '), mean_pm(r['linear_probe_accuracy_mean'], r['linear_probe_accuracy_std']), f"{r['silhouette_mean']:.3f} ± {r['silhouette_std']:.3f}", f"{r['fisher_ratio_mean']:.3f} ± {r['fisher_ratio_std']:.3f}", mean_pm(r['activation_sparsity_mean'],r['activation_sparsity_std']), mean_pm(r['dead_unit_rate_mean'],r['dead_unit_rate_std'])])
add_table(doc,['Representation','Linear probe','Silhouette','Fisher ratio','Sparsity','Dead-unit rate'],rows,'Table 6. Digits representation geometry and feature health. Source: representation_geometry_summary.csv.')
doc.add_paragraph(
"On digits, raw pixels and Oja features have positive silhouette scores and higher Fisher ratios than Sparse Oja variants. Sparse and Homeostatic Sparse Oja show negative silhouette scores, high dead-unit rates, and weak class selectivity despite high sparsity. This supports the mechanism-level claim that sparsity alone does not guarantee class-aligned geometry.")
for img,cap in [('representation_geometry_digits_summary.png','Figure 6. Representation geometry and feature-health summary on digits. Source: representation_geometry_summary.csv.'),('representation_pca_digits.png','Figure 7. PCA visualizations of digit representation spaces. Source: run_representation_geometry.py.')]:
    if (FIG/img).exists(): doc.add_picture(str(FIG/img), width=Inches(6.2)); add_caption(doc, cap)

doc.add_page_break()
doc.add_heading('E. Continual learning and replay', level=2)
# Condense 0 and 20 percent replay for MNIST and Fashion-MNIST
sub = replay[replay['replay_frac'].isin([0.0,0.2]) & replay['dataset'].isin(['mnist','fashion_mnist'])]
rows=[]
for _,r in sub.iterrows():
    rows.append([r['dataset'].replace('_','-'), r['model'], f"{100*r['replay_frac']:.0f}%", mean_pm(r['A_before_mean'],r['A_before_std']), mean_pm(r['A_after_mean'],r['A_after_std']), mean_pm(r['A_T2_mean'],r['A_T2_std']), f"{r['normalized_forgetting_mean']:.2f} ± {r['normalized_forgetting_std']:.2f}"])
add_table(doc,['Dataset','Model','Replay','A_before','A_after','A_T2','Norm. forgetting'],rows,'Table 7. Replay sweep subset results, condensed to 0% and 20% replay. Source: replay_sweep_subset_summary.csv.')
if (FIG/'replay_sweep_partial_mnist.png').exists(): doc.add_picture(str(FIG/'replay_sweep_partial_mnist.png'), width=Inches(5.8)); add_caption(doc,'Figure 8. Replay sweep on MNIST subset. Source: replay_sweep_subset_summary.csv.')

doc.add_page_break()
doc.add_heading('F. Statistical robustness', level=2)
rows=[]
for _,r in stat.head(8).iterrows():
    rows.append([r['dataset'].replace('_','-'), f"{r['model_a']} vs {r['model_b']}", f"{r['mean_diff_a_minus_b']:.3f}", f"{r['wilcoxon_p']:.4f}", f"{r['cohens_dz']:.2f}"])
add_table(doc,['Dataset','Comparison','Mean diff','Wilcoxon p','Cohen dz'],rows,'Table 8. Paired statistical tests across seeds. With five seeds, p-values are supportive only; effect sizes and cross-dataset consistency are emphasized. Source: statistical_tests_summary.csv.')

if (FIG/'mechanism_map_v2.png').exists(): doc.add_picture(str(FIG/'mechanism_map_v2.png'), width=Inches(6.2)); add_caption(doc,'Figure 9. Mechanism map summarizing the central interpretation. Source: make_enhanced_figures.py.')

# Discussion
doc.add_heading('7. Discussion', level=1)
doc.add_paragraph(
"The upgraded study strengthens the negative result. Supervised models and the Standard CNN dominate task-discriminative accuracy, while simple Oja/Sparse Oja remain weak on MNIST and Fashion-MNIST. Homeostatic Sparse Oja improves over fixed Sparse Oja in some cases, especially digits and Fashion-MNIST, but it does not close the gap to backpropagation. The sparsity sweep is the clearest mechanism-isolation result: when active units are extremely limited, representations are sparse but weakly discriminative; as active fraction increases, accuracy improves while sparsity declines. This separates two often-conflated properties: sparse activity and useful class separation.")
doc.add_paragraph(
"The new frozen-readout and geometry analyses refine the interpretation. On digits, Oja features become highly accurate with a stronger readout, suggesting that some local features are not inherently useless. However, Sparse Oja and Homeostatic Sparse Oja remain much weaker even with linear SVM, KNN, and other readouts, and their geometry metrics show low or negative class separation. This supports a nuanced conclusion: local Oja features can preserve information on simple data, but extreme sparsity and homeostatic thresholds can reduce linearly or geometrically accessible task structure.")
doc.add_paragraph(
"The replay analysis shows that memory stabilization changes forgetting, but replay alone does not make local-feature models competitive. Crucially, forgetting is only meaningful when Task 1 is first learned well. Thus, the study does not claim that Hebbian/Oja learning beats backpropagation. Instead, it argues that biological efficiency likely emerges from interacting mechanisms: local plasticity, sparse coding, competition, homeostasis, replay, feedback, neuromodulation, recurrence, and task-relevant signals.")

# Limitations
doc.add_heading('8. Limitations', level=1)
for item in [
"The submitted reproducibility ZIP does not include raw MNIST/Fashion-MNIST IDX files, so new frozen-readout and representation-geometry reruns for those datasets are blocked in this pass. Existing MNIST/Fashion-MNIST benchmark CSVs are retained from the completed prior package.",
"The Standard CNN is a credible LeNet-style baseline but is not a state-of-the-art architecture.",
"Homeostatic Sparse Oja uses a simple adaptive-threshold mechanism and does not include neuromodulation, recurrent feedback, dendritic compartments, or reward-modulated plasticity.",
"Sparsity and replay sweeps are fixed-subset mechanism-isolation analyses, not full-data state-of-the-art benchmarks.",
"Only five seeds are available for most full benchmark results. Digits frozen-readout results include five complete local-feature seeds; some supervised comparison feature sources completed only two seeds and are labeled as such.",
"Approximate MACs are computational proxies, not direct energy measurements.",
"No test-set tuning was performed, but hyperparameters were not exhaustively optimized." ]:
    doc.add_paragraph(item, style='List Bullet')

# Conclusion
doc.add_heading('9. Conclusion', level=1)
doc.add_paragraph(
"SparseMechanismBench turns the project from a simple comparison of Oja versus backpropagation into a mechanism-isolation framework. Across completed benchmarks, backpropagation-trained and supervised models remain stronger for classification accuracy and sample efficiency. Sparse Oja-style learning reliably creates sparse representations, but sparse local plasticity alone is insufficient for task-discriminative learning or stable continual learning. Homeostatic competition and replay improve some trade-offs but do not close the gap. The strongest conclusion is therefore not that local plasticity is worse in principle, but that biological learning efficiency likely requires coordinated mechanisms: local plasticity, sparse coding, competition, homeostasis, replay, feedback, and task-relevant signals.")

# Code data
doc.add_heading('10. Code and Data Availability', level=1)
doc.add_paragraph(
"The accompanying reproducibility package contains README.md, requirements.txt, data_loading.py, models.py, train_classification.py, train_hebbian_oja.py, train_sparse_oja.py, train_homeostatic_oja.py, run_sample_efficiency.py, run_continual_learning.py, run_sparsity_sweep_subset.py, run_replay_sweep_subset.py, run_frozen_readout.py, run_representation_geometry.py, run_full_benchmark.py, plot_results.py, make_enhanced_figures.py, raw CSV files, summary CSV files, generated figures, and scripts. A public GitHub/OSF/Zenodo link should be added at submission time. All reported numerical results must be traceable to raw CSV files and plotting scripts.")

# One-page summary appendix
doc.add_heading('Appendix A. STS One-Page Summary', level=1)
doc.add_paragraph((ROOT/'one_page_summary.txt').read_text())

doc.add_heading('Appendix B. Application Descriptions', level=1)
doc.add_paragraph((ROOT/'application_descriptions.txt').read_text())

doc.add_heading('Appendix C. Reviewer-Response Memo', level=1)
doc.add_paragraph((ROOT/'reviewer_response_memo.txt').read_text())

# References
doc.add_heading('References', level=1)
refs=[
"Hebb, D. O. (1949). The organization of behavior: A neuropsychological theory. Wiley.",
"Oja, E. (1982). A simplified neuron model as a principal component analyzer. Journal of Mathematical Biology, 15(3), 267-273. https://doi.org/10.1007/BF00275687",
"Bi, G.-Q., & Poo, M.-M. (1998). Synaptic modifications in cultured hippocampal neurons: Dependence on spike timing, synaptic strength, and postsynaptic cell type. The Journal of Neuroscience, 18(24), 10464-10472. https://doi.org/10.1523/JNEUROSCI.18-24-10464.1998",
"Bienenstock, E. L., Cooper, L. N., & Munro, P. W. (1982). Theory for the development of neuron selectivity: Orientation specificity and binocular interaction in visual cortex. The Journal of Neuroscience, 2(1), 32-48. https://doi.org/10.1523/JNEUROSCI.02-01-00032.1982",
"Rao, R. P. N., & Ballard, D. H. (1999). Predictive coding in the visual cortex: A functional interpretation of some extra-classical receptive-field effects. Nature Neuroscience, 2, 79-87. https://doi.org/10.1038/4580",
"Olshausen, B. A., & Field, D. J. (1996). Emergence of simple-cell receptive field properties by learning a sparse code for natural images. Nature, 381, 607-609. https://doi.org/10.1038/381607a0",
"Olshausen, B. A., & Field, D. J. (1997). Sparse coding with an overcomplete basis set: A strategy employed by V1? Vision Research, 37(23), 3311-3325. https://doi.org/10.1016/S0042-6989(97)00169-7",
"Makhzani, A., & Frey, B. (2013). k-sparse autoencoders. arXiv:1312.5663. https://arxiv.org/abs/1312.5663",
"Rumelhart, D. E., Hinton, G. E., & Williams, R. J. (1986). Learning representations by back-propagating errors. Nature, 323, 533-536. https://doi.org/10.1038/323533a0",
"Lillicrap, T. P., Cownden, D., Tweed, D. B., & Akerman, C. J. (2016). Random synaptic feedback weights support error backpropagation for deep learning. Nature Communications, 7, 13276. https://doi.org/10.1038/ncomms13276",
"Scellier, B., & Bengio, Y. (2017). Equilibrium propagation: Bridging the gap between energy-based models and backpropagation. Frontiers in Computational Neuroscience, 11, 24. https://doi.org/10.3389/fncom.2017.00024",
"Lee, D.-H., Zhang, S., Fischer, A., & Bengio, Y. (2015). Difference target propagation. In Machine Learning and Knowledge Discovery in Databases (ECML PKDD 2015). https://arxiv.org/abs/1412.7525",
"Whittington, J. C. R., & Bogacz, R. (2019). Theories of error back-propagation in the brain. Trends in Cognitive Sciences, 23(3), 235-250. https://doi.org/10.1016/j.tics.2018.12.005",
"McCloskey, M., & Cohen, N. J. (1989). Catastrophic interference in connectionist networks: The sequential learning problem. In G. H. Bower (Ed.), Psychology of Learning and Motivation (Vol. 24, pp. 109-165). Academic Press. https://doi.org/10.1016/S0079-7421(08)60536-8",
"McClelland, J. L., McNaughton, B. L., & O'Reilly, R. C. (1995). Why there are complementary learning systems in the hippocampus and neocortex. Psychological Review, 102(3), 419-457. https://doi.org/10.1037/0033-295X.102.3.419",
"Lin, L.-J. (1992). Self-improving reactive agents based on reinforcement learning, planning and teaching. Machine Learning, 8, 293-321. https://doi.org/10.1007/BF00992699",
"Mnih, V., Kavukcuoglu, K., Silver, D., et al. (2015). Human-level control through deep reinforcement learning. Nature, 518, 529-533. https://doi.org/10.1038/nature14236",
"Kirkpatrick, J., Pascanu, R., Rabinowitz, N., et al. (2017). Overcoming catastrophic forgetting in neural networks. Proceedings of the National Academy of Sciences, 114(13), 3521-3526. https://doi.org/10.1073/pnas.1611835114",
"LeCun, Y., Bottou, L., Bengio, Y., & Haffner, P. (1998). Gradient-based learning applied to document recognition. Proceedings of the IEEE, 86(11), 2278-2324. https://doi.org/10.1109/5.726791",
"Xiao, H., Rasul, K., & Vollgraf, R. (2017). Fashion-MNIST: A novel image dataset for benchmarking machine learning algorithms. arXiv:1708.07747. https://arxiv.org/abs/1708.07747" ]
for ref in refs:
    p=doc.add_paragraph(ref)
    p.paragraph_format.space_after=Pt(3)

# Save
OUT.parent.mkdir(exist_ok=True)
doc.save(OUT)
print('wrote',OUT)
