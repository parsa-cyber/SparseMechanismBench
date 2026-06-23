from pathlib import Path
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT=Path(__file__).resolve().parent
RESULTS=ROOT/'results'
FIG=ROOT/'figures'
OUT=Path('/mnt/data/SparseMechanismBench_final')
OUT.mkdir(exist_ok=True)

def pct(x): return f"{100*x:.1f}%"
def pm(mean,std): return f"{100*mean:.1f}% ± {100*std:.1f}%"

def set_cell_shading(cell, fill='D9EAF7'):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)

def style_table(table):
    table.style='Table Grid'; table.alignment=WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                for run in p.runs: run.font.size=Pt(8)
    for cell in table.rows[0].cells:
        set_cell_shading(cell)
        for p in cell.paragraphs:
            for run in p.runs: run.bold=True

def add_table(doc, headers, rows, caption):
    table=doc.add_table(rows=1, cols=len(headers)); style_table(table)
    for i,h in enumerate(headers): table.rows[0].cells[i].text=str(h)
    for r in rows:
        cells=table.add_row().cells
        for i,v in enumerate(r): cells[i].text=str(v)
    p=doc.add_paragraph(caption); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs: run.font.size=Pt(9); run.italic=True
    return table

def add_fig(doc, path, caption, width=6.3):
    if Path(path).exists():
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(width))
        cap=doc.add_paragraph(caption); cap.alignment=WD_ALIGN_PARAGRAPH.CENTER
        for run in cap.runs: run.font.size=Pt(9); run.italic=True
    else:
        doc.add_paragraph(f"[Figure missing: {path}]")

# Load results
class_df=pd.read_csv(RESULTS/'updated_classification_summary.csv')
readout=pd.read_csv(RESULTS/'frozen_readout_mnist_fashion_summary.csv')
geom=pd.read_csv(RESULTS/'representation_geometry_mnist_fashion_summary.csv')
sweep=pd.read_csv(RESULTS/'sparsity_sweep_subset_summary.csv')
replay=pd.read_csv(RESULTS/'replay_sweep_subset_summary.csv')
stats=pd.read_csv(RESULTS/'frozen_readout_statistical_tests.csv')
selectivity=pd.read_csv(RESULTS/'selectivity_summary.csv')

# Begin doc
doc=Document()
sec=doc.sections[0]
sec.top_margin=Inches(0.7); sec.bottom_margin=Inches(0.7); sec.left_margin=Inches(0.75); sec.right_margin=Inches(0.75)
styles=doc.styles
styles['Normal'].font.name='Times New Roman'; styles['Normal'].font.size=Pt(10.5)
styles['Heading 1'].font.name='Arial'; styles['Heading 1'].font.size=Pt(16)
styles['Heading 2'].font.name='Arial'; styles['Heading 2'].font.size=Pt(13)
styles['Heading 3'].font.name='Arial'; styles['Heading 3'].font.size=Pt(11)

title='SparseMechanismBench: Why Sparse Local Plasticity Is Not Enough for Task-Discriminative Learning'
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run(title); r.bold=True; r.font.size=Pt(18); r.font.name='Arial'
p=doc.add_paragraph('Parsa Fatehi / Langley High School'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER

# Abstract
doc.add_heading('Abstract',1)
doc.add_paragraph(
"Biological neural systems rely heavily on local plasticity, sparse activity, competition, homeostasis, replay, and feedback, whereas modern artificial networks usually rely on global gradient-based optimization. This study introduces SparseMechanismBench, a mechanism-isolation framework testing whether sparse local Hebbian/Oja-style plasticity is sufficient for task-discriminative learning and continual-learning stability. Across digits, MNIST, and Fashion-MNIST, prior benchmark runs showed that supervised and backpropagation-based models strongly outperform simple Oja and Sparse Oja feature learners in classification accuracy. I extend the work with a LeNet-style CNN baseline, Homeostatic Sparse Oja, sparsity and replay sweeps, frozen-feature readout experiments, representation-geometry diagnostics, feature-health metrics, visualization, and statistical comparisons. The new MNIST/Fashion-MNIST frozen-readout experiments used fixed 2,000-train/1,000-test mechanism-isolation subsets and 10 seeds for local and raw-pixel features. Results show that Oja features can contain task-relevant information that is accessible to a ridge readout, but extreme Sparse Oja and Homeostatic Sparse Oja produce highly sparse representations that are much less geometrically class-separated. On MNIST, ridge readout accuracy was 80.7% ± 1.6% for Oja but only 31.7% ± 5.2% for Sparse Oja and 32.9% ± 5.5% for Homeostatic Sparse Oja. On Fashion-MNIST, Oja reached 76.7% ± 1.5%, while Sparse Oja and Homeostatic Sparse Oja reached about 60%. These results preserve and sharpen the central negative result: sparsity alone is not enough. Useful biological-style learning likely requires interacting mechanisms, including local plasticity, sparse coding, homeostasis, replay, and task-relevant credit signals." )

# Intro
doc.add_heading('1. Introduction',1)
doc.add_paragraph(
"A central puzzle in neuroscience and machine learning is how brains learn efficiently from local synaptic updates while modern artificial neural networks rely on global backpropagation. Hebbian and Oja-style learning rules are biologically plausible because they use information locally available at the synapse, but they do not directly optimize a supervised task objective. Sparse coding is often proposed as a source of biological efficiency, yet sparse representations are not automatically class-separable or useful for a downstream task. This paper therefore asks a sharper mechanism-level question: which biological constraints are sufficient, and which are insufficient, for closing the gap between local plasticity and gradient-based optimization?" )
doc.add_paragraph(
"SparseMechanismBench treats local plasticity, sparsity, competition, homeostasis, replay, and supervised readout strength as separable experimental factors. The goal is not to show that Hebbian/Oja learning beats backpropagation. The goal is to test whether biologically inspired mechanisms improve the accuracy-sparsity-forgetting trade-off, and to preserve negative results when they show that an intuitive mechanism is insufficient.")

doc.add_heading('Scientific Contribution',2)
doc.add_paragraph(
"This study makes five contributions. First, it separates representational sparsity from task-discriminative accuracy by measuring both under the same pipeline. Second, it compares local and gradient-based learning systems across digits, MNIST, and Fashion-MNIST. Third, it tests a mechanism ladder: Oja, Sparse Oja, Homeostatic Sparse Oja, replay, MLPs, and standard CNNs. Fourth, it adds a frozen-feature readout experiment to determine whether local features fail because the representation is poor or because the readout is too weak. Fifth, it introduces representation-geometry and feature-health diagnostics that quantify whether sparse local representations are class-separated, utilized, and selective. The main scientific result is that sparse local plasticity creates sparse representations, but by itself it does not reliably create task-discriminative or stable representations.")

# Related work
doc.add_heading('2. Related Work',1)
doc.add_heading('2.1 Biological local learning',2)
doc.add_paragraph(
"Hebb (1949) proposed that synapses strengthen when pre- and postsynaptic neurons are co-active, establishing the classic local-learning principle often summarized as 'cells that fire together wire together.' Oja (1982) added a stabilizing normalization term, making Hebbian learning behave like a principal-component analyzer under simple assumptions. Spike-timing-dependent plasticity further refines local learning by linking synaptic updates to the precise temporal order of pre- and postsynaptic spikes (Bi & Poo, 1998). BCM theory introduced activity-dependent thresholds that regulate plasticity and prevent runaway firing (Bienenstock, Cooper, & Munro, 1982). Predictive-coding theories suggest that local error signals and recurrent dynamics may approximate gradient-like learning while remaining more biologically plausible than literal backpropagation (Rao & Ballard, 1999; Whittington & Bogacz, 2019). These mechanisms are biologically plausible because they depend on local or semi-local variables, but they also raise the central credit-assignment problem: local correlations alone do not necessarily indicate which synapses are responsible for task success." )
doc.add_heading('2.2 Sparse coding and efficient representations',2)
doc.add_paragraph(
"Sparse coding is a major theory of efficient neural representation. Olshausen and Field (1996) showed that sparse coding of natural images can yield receptive fields resembling simple cells in visual cortex. Sparse autoencoders and k-winners-take-all mechanisms similarly encourage representations in which only a subset of units is active at once. Sparse activity can reduce metabolic cost, decorrelate representations, and improve interpretability. However, sparsity is a representational constraint, not a supervised objective. A sparse code can be efficient but still poorly aligned with class labels. This distinction is central to SparseMechanismBench: the experiments explicitly test whether sparse Oja-style features become more class-discriminative, or whether they merely become sparse." )
doc.add_heading('2.3 Biologically plausible credit assignment',2)
doc.add_paragraph(
"Backpropagation assigns credit by differentiating a global loss through every layer (Rumelhart, Hinton, & Williams, 1986). This is effective but biologically controversial because it seems to require nonlocal error signals, exact weight transport, and coordinated layerwise updates. Feedback alignment relaxes the need for symmetric feedback weights and shows that random feedback can sometimes support learning (Lillicrap et al., 2016). Equilibrium propagation (Scellier & Bengio, 2017), target propagation (Bengio, 2014), predictive coding (Whittington & Bogacz, 2019), and dendritic-error models all attempt to make credit assignment more biologically plausible. These approaches are directly relevant because simple Oja and Sparse Oja lack task-directed credit signals. If local features fail even with stronger readouts, the missing component may be representation quality; if nonlinear readouts recover performance, the local representation may contain task information that is not linearly accessible." )
doc.add_heading('2.4 Catastrophic forgetting and replay',2)
doc.add_paragraph(
"Catastrophic forgetting occurs when learning a new task overwrites representations needed for an earlier task (McCloskey & Cohen, 1989). Complementary learning systems theory proposes that the hippocampus and neocortex support rapid episodic encoding and slower integrated learning (McClelland, McNaughton, & O'Reilly, 1995). Modern continual-learning methods include experience replay, generative replay, and regularization methods such as elastic weight consolidation (Kirkpatrick et al., 2017). Replay is biologically motivated by hippocampal reactivation, but forgetting metrics are only meaningful when Task 1 is first learned well. A model that never learns Task 1 can appear to forget little. This paper therefore reports Task 1 accuracy before Task 2, Task 1 accuracy after Task 2, Task 2 accuracy, average final accuracy, raw forgetting, and normalized forgetting." )
doc.add_heading('2.5 Why this work is different',2)
doc.add_paragraph(
"Much prior work studies one mechanism at a time: Hebbian/Oja learning, sparse coding, biologically plausible credit assignment, or replay. SparseMechanismBench instead tests these mechanisms in one reproducible pipeline. The design separates local plasticity, sparsity, competition/homeostasis, replay, readout strength, representation geometry, and supervised baselines. This enables a more specific conclusion than 'Hebbian learning loses to backpropagation': sparse local plasticity alone can create sparse codes, but it does not by itself guarantee class separation, task-accessible information, or retention." )

# Research questions
doc.add_heading('3. Research Question and Hypotheses',1)
doc.add_paragraph("Core question: are sparse local-plasticity representations weak because the representations themselves are poorly task-aligned, or because the downstream readout/classifier is too weak?")
for h in [
"H1: Supervised and backpropagation-based models will outperform simple Oja/Sparse Oja feature learners in classification accuracy and sample efficiency.",
"H2: Sparse Oja and Homeostatic Sparse Oja will produce much higher activation sparsity than MLP and CNN features.",
"H3: If stronger frozen readouts still fail on Sparse Oja features, then the representation itself is poorly task-aligned; if stronger readouts substantially improve accuracy, then the features contain information but are not accessible to a weak readout.",
"H4: Replay will reduce raw forgetting only when Task 1 is learned well enough for forgetting to be meaningful.",
"H5: Homeostasis and competition may improve the accuracy-sparsity trade-off, but they are not expected to fully replace task-relevant credit assignment."
]: doc.add_paragraph(h, style='List Bullet')

# Mathematical setup
doc.add_heading('4. Mathematical Setup',1)
doc.add_paragraph("Hebbian learning updates a synapse from presynaptic unit j to postsynaptic unit i using local activities:")
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run("Δwᵢⱼ = η xⱼ yᵢ").bold=True
doc.add_paragraph("where η is the learning rate, xⱼ is presynaptic activity, and yᵢ is postsynaptic activity.")
doc.add_paragraph("Oja's rule adds an activity-dependent stabilizing term:")
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run("Δwᵢ = η (yᵢ x − yᵢ² wᵢ)").bold=True
doc.add_paragraph("Backpropagation instead updates weights using the gradient of a loss L:")
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run("w ← w − η ∂L/∂w").bold=True
doc.add_paragraph("Forgetting is measured as:")
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run("F = A_before − A_after").bold=True
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run("F_norm = (A_before − A_after) / A_before").bold=True

# Methods
doc.add_heading('5. Methods',1)
doc.add_heading('5.1 Datasets and validation',2)
doc.add_paragraph("Experiments use scikit-learn digits, MNIST IDX, and Fashion-MNIST IDX. Pixels are normalized to [0, 1]. No test-set tuning was performed. Full benchmark results use seeds 0-4; the new frozen-readout and representation-geometry experiments use fixed 2,000-train/1,000-test mechanism-isolation subsets, with 10 seeds for local and raw-pixel features and 5 seeds for supervised hidden-feature baselines.")
# dataset table
val=pd.read_csv(RESULTS/'dataset_validation.csv')
rows=[[r['dataset'], int(r['train_size']), int(r['test_size']), r['image_shape'], f"{r['train_min']:.1f}-{r['train_max']:.1f}", f"{r['test_min']:.1f}-{r['test_max']:.1f}"] for _,r in val.iterrows()]
add_table(doc,['Dataset','Train','Test','Shape','Train range','Test range'],rows,'Table 1. Dataset validation. Source: results/dataset_validation.csv.')
add_fig(doc, FIG/'example_images_grid.png','Figure 1. Example images from digits, MNIST, and Fashion-MNIST. Source: figures/example_images_grid.png.', width=5.7)

doc.add_heading('5.2 Models',2)
doc.add_paragraph("The benchmark compares logistic regression, random features, PCA features, MLP, MLP+dropout, compact runtime-limited CNN, LeNet-style Standard CNN, Oja, Sparse Oja, and Homeostatic Sparse Oja. Sparse Oja uses k-winners-take-all competition. Homeostatic Sparse Oja adds adaptive thresholds: neurons firing above the target rate increase threshold, and neurons firing below target decrease threshold. Local-feature models are unsupervised at the feature stage, after which supervised readouts are trained on frozen features.")
mechanism_rows=[['Oja','Yes','No','No','Optional','No'],['Sparse Oja','Yes','Yes (k-WTA)','No','Optional','No'],['Homeostatic Sparse Oja','Yes','Yes','Yes','Optional','No'],['MLP/CNN','No','No','No','Optional','Yes'],['Frozen readouts','No','N/A','N/A','No','Yes readout only']]
add_table(doc,['Model','Local?','Sparse?','Homeostatic?','Replay?','Task-directed signal?'],mechanism_rows,'Table 2. Mechanism ladder tested in SparseMechanismBench.')

doc.add_heading('5.3 Frozen Feature Readout Experiment',2)
doc.add_paragraph("For Oja, Sparse Oja, and Homeostatic Sparse Oja, features are learned unsupervised on the training subset only and then frozen. Multiple supervised readouts are trained on the frozen features: SGD logistic regression, ridge classifier, linear SVM-style SGD hinge classifier, k-nearest neighbors, and a small MLP readout. Controls include raw pixels, MLP hidden features, and Standard CNN penultimate features. This experiment answers whether local features fail because the feature representation is weak or because the readout is too limited.")

doc.add_heading('5.4 Representation Geometry and Feature Health',2)
doc.add_paragraph("For each representation, I measured linear-probe accuracy, silhouette score, Davies-Bouldin index, Fisher between/within scatter ratio, mean intra-class distance, mean inter-class distance, inter/intra distance ratio, activation sparsity, dead-unit rate, mean activation, utilization entropy, and class selectivity. PCA visualizations, learned-feature grids, activation histograms, and confusion matrices provide qualitative diagnostics.")

doc.add_heading('5.5 Continual Learning, Replay, and Statistics',2)
doc.add_paragraph("Task 1 contains classes 0-4 and Task 2 contains classes 5-9. Replay sweeps test replay rates of 0%, 1%, 5%, 10%, and 20% using fixed mechanism-isolation subsets. Statistical analysis reports means, standard deviations, 95% confidence intervals, paired tests, Wilcoxon tests, and Cohen's dz. With small seed counts, p-values are treated as supportive, not definitive.")

# Results
doc.add_heading('6. Results',1)
doc.add_heading('6.1 Supervised baselines vs. local plasticity',2)
# classification table concise for MNIST/Fashion + digits maybe
models_order=['logistic_regression','mlp','mlp_dropout','compact_cnn','standard_cnn','hebbian_oja','sparse_hebbian_oja','homeostatic_sparse_oja']
rows=[]
for m in models_order:
    r=[]
    display={'logistic_regression':'LogReg','mlp':'MLP','mlp_dropout':'MLP+dropout','compact_cnn':'compact CNN','standard_cnn':'Standard CNN','hebbian_oja':'Oja','sparse_hebbian_oja':'Sparse Oja','homeostatic_sparse_oja':'Homeostatic Sparse Oja'}[m]
    r.append(display)
    for ds in ['digits','mnist','fashion_mnist']:
        sub=class_df[(class_df.dataset==ds)&(class_df.model==m)]
        if len(sub): r.append(pm(sub.accuracy_mean.iloc[0], sub.accuracy_std.iloc[0]))
        else: r.append('N/A')
    rows.append(r)
add_table(doc,['Model','digits','MNIST','Fashion-MNIST'],rows,'Table 3. Full benchmark classification accuracy. Source: results/updated_classification_summary.csv.')
add_fig(doc, FIG/'standard_vs_compact_cnn.png','Figure 2. Standard CNN removes the weak-CNN concern by outperforming compact CNN on MNIST/Fashion-MNIST. Source: figures/standard_vs_compact_cnn.png.',width=5.6)
add_fig(doc, FIG/'mechanism_ladder_accuracy.png','Figure 3. Mechanism-ladder classification results. Source: figures/mechanism_ladder_accuracy.png.',width=6.3)

doc.add_heading('6.2 Sparsity-accuracy trade-off',2)
doc.add_paragraph("The sparsity sweep tests whether Sparse Oja failed because 90% sparsity was simply too extreme. On MNIST, accuracy improved as more units became active: 35.1% at 95.3% sparsity, 62.6% at 79.7% sparsity, and 79.7% at about 18.8% sparsity. Fashion-MNIST showed the same general pattern, with best accuracy at the least sparse setting. This supports the claim that extreme sparsity can improve efficiency while damaging task discrimination.")
rows=[]
for _,r in sweep.iterrows():
    if r['model']=='Sparse Oja':
        rows.append([r['dataset'].replace('_','-'),f"{int(r['active_frac']*100)}%",pm(r['sparsity_mean'],r['sparsity_std']),pm(r['accuracy_mean'],r['accuracy_std'])])
add_table(doc,['Dataset','Active units','Sparsity','Accuracy'],rows,'Table 4. Sparse Oja sparsity sweep on fixed subsets. Source: results/sparsity_sweep_subset_summary.csv.')
add_fig(doc, FIG/'ablation_accuracy_vs_sparsity.png','Figure 4. Accuracy-sparsity sweep: extreme sparsity weakens task discrimination. Source: figures/ablation_accuracy_vs_sparsity.png.',width=6.3)

doc.add_heading('6.3 Frozen readout: feature quality vs. readout weakness',2)
doc.add_paragraph("The frozen-readout experiment is the key new diagnostic. If a strong readout recovers high accuracy, the representation contains task information but may not be linearly accessible. If strong readouts still fail, the representation is poorly task-aligned. On MNIST and Fashion-MNIST, Oja features are substantially more readable than Sparse Oja features. Ridge readout recovered 80.7% ± 1.6% on MNIST and 76.7% ± 1.5% on Fashion-MNIST from Oja features. In contrast, 90%-sparse features remained far weaker: Sparse Oja reached 31.7% ± 5.2% on MNIST and 60.1% ± 3.5% on Fashion-MNIST with ridge readout. Homeostatic Sparse Oja was similar or slightly better/worse depending on dataset, but did not close the gap.")
# best readout per dataset/source table
rows=[]
for ds in ['mnist','fashion_mnist']:
    for source in ['raw_pixels','oja','sparse_oja','homeostatic_sparse_oja','mlp_hidden','standard_cnn_penultimate']:
        sub=readout[(readout.dataset==ds)&(readout.feature_source==source)]
        if sub.empty: continue
        best=sub.sort_values('accuracy_mean',ascending=False).iloc[0]
        rows.append([ds.replace('_','-'), source.replace('_',' '), best['readout'], pm(best['accuracy_mean'], best['accuracy_std']), int(best['n']), pm(best['feature_sparsity_mean'], best['feature_sparsity_std'])])
add_table(doc,['Dataset','Feature source','Best readout','Accuracy','n','Feature sparsity'],rows,'Table 5. Frozen-feature readout results on fixed 2,000-train/1,000-test subsets. Source: results/frozen_readout_mnist_fashion_summary.csv.')
add_fig(doc, FIG/'frozen_readout_mnist_heatmap.png','Figure 5. MNIST frozen-readout heatmap. Source: figures/frozen_readout_mnist_heatmap.png and results/frozen_readout_mnist_fashion_summary.csv.',width=6.2)
add_fig(doc, FIG/'frozen_readout_fashion_mnist_heatmap.png','Figure 6. Fashion-MNIST frozen-readout heatmap. Source: figures/frozen_readout_fashion_mnist_heatmap.png and results/frozen_readout_mnist_fashion_summary.csv.',width=6.2)

doc.add_heading('6.4 Representation geometry',2)
doc.add_paragraph("Representation-geometry metrics show that sparse local features are not merely lower-performing because of one weak classifier. Their class geometry is also poorer. On MNIST, Sparse Oja and Homeostatic Sparse Oja have negative silhouette scores and much lower inter/intra ratios than MLP or CNN features. The dead-unit and utilization metrics show that k-WTA sparse features are consistently active at the target sparsity level, so the issue is not simply inactive code; it is weak class alignment.")
rows=[]
for _,r in geom.iterrows():
    rows.append([r['dataset'].replace('_','-'), r['feature_source'].replace('_',' '), pm(r['linear_probe_accuracy_mean'],r['linear_probe_accuracy_std']), f"{r['silhouette_mean']:.3f}", f"{r['inter_intra_ratio_mean']:.3f}", pm(r['activation_sparsity_mean'],r['activation_sparsity_std']), f"{r['class_selectivity_index_mean']:.3f}"])
add_table(doc,['Dataset','Feature source','Linear probe','Silhouette','Inter/intra','Sparsity','Selectivity'],rows,'Table 6. Representation geometry and feature-health summary. Source: results/representation_geometry_mnist_fashion_summary.csv.')
add_fig(doc, FIG/'representation_geometry_mnist_summary.png','Figure 7. MNIST representation-geometry summary. Source: figures/representation_geometry_mnist_summary.png.',width=6.3)
add_fig(doc, FIG/'representation_geometry_fashion_mnist_summary.png','Figure 8. Fashion-MNIST representation-geometry summary. Source: figures/representation_geometry_fashion_mnist_summary.png.',width=6.3)
add_fig(doc, FIG/'representation_pca_mnist_uploaded.png','Figure 9. PCA visualization of MNIST representations, seed 0. Source: figures/representation_pca_mnist_uploaded.png.',width=6.3)
add_fig(doc, FIG/'representation_pca_fashion_mnist_uploaded.png','Figure 10. PCA visualization of Fashion-MNIST representations, seed 0. Source: figures/representation_pca_fashion_mnist_uploaded.png.',width=6.3)

doc.add_heading('6.5 Continual learning and replay',2)
doc.add_paragraph("Replay reduces forgetting but does not make local-feature models competitive. In no-replay conditions, Task 1 accuracy after Task 2 often collapses to 0%, producing normalized forgetting near 1.0. With 20% replay, Oja on MNIST retains substantially more Task 1 performance than without replay, but still loses much of its earlier knowledge. These results support the interpretation that replay is a stabilizing mechanism, not a complete substitute for task-directed learning.")
rows=[]
for ds in ['mnist','fashion_mnist']:
    for model in ['oja','sparse_oja','homeostatic_sparse_oja']:
        for rf in [0.0,0.2]:
            sub=replay[(replay.dataset==ds)&(replay.model==model)&(replay.replay_frac==rf)]
            if not sub.empty:
                r=sub.iloc[0]
                rows.append([ds.replace('_','-'),model.replace('_',' '),f"{int(rf*100)}%",pm(r['A_before_mean'],r['A_before_std']),pm(r['A_after_mean'],r['A_after_std']),pm(r['A_T2_mean'],r['A_T2_std']),pm(r['normalized_forgetting_mean'],r['normalized_forgetting_std'])])
add_table(doc,['Dataset','Model','Replay','A_before','A_after','A_T2','Norm. forgetting'],rows,'Table 7. Replay sweep condensed to 0% and 20% replay. Full results in results/replay_sweep_subset_summary.csv.')
add_fig(doc, FIG/'replay_sweep_partial_mnist.png','Figure 11. MNIST replay sweep for Oja/Sparse Oja/Homeostatic Sparse Oja. Source: figures/replay_sweep_partial_mnist.png.',width=5.8)

doc.add_heading('6.6 Interpretability diagnostics',2)
doc.add_paragraph("Learned feature grids and activation histograms show that local features are structured but not necessarily class-aligned. Confusion matrices from ridge readouts further show which classes collapse under local representations. These figures support the central distinction between sparse representation and task-discriminative representation.")
add_fig(doc, FIG/'learned_features_mnist.png','Figure 12. Learned MNIST local features for Oja, Sparse Oja, and Homeostatic Sparse Oja. Source: figures/learned_features_mnist.png.',width=6.3)
add_fig(doc, FIG/'confusion_matrices_mnist.png','Figure 13. MNIST confusion matrices using ridge readouts, seed 0. Source: figures/confusion_matrices_mnist.png.',width=6.3)
add_fig(doc, FIG/'activation_histograms_mnist.png','Figure 14. MNIST activation histograms, seed 0. Source: figures/activation_histograms_mnist.png.',width=6.3)

doc.add_heading('6.7 Statistical robustness',2)
doc.add_paragraph("Statistical tests use paired seeds whenever possible. With only five to ten seeds, p-values are treated as supportive rather than definitive. The most robust pattern is effect size and cross-dataset consistency: Sparse Oja is consistently more sparse but much less task-discriminative than raw pixels, Oja, MLP, or CNN features.")
rows=[]
for _,r in stats.iterrows():
    if r['comparison'] in ['oja vs sparse_oja','sparse_oja vs homeostatic_sparse_oja','raw_pixels vs sparse_oja','raw_pixels vs oja','mlp_hidden vs oja','standard_cnn_penultimate vs oja']:
        rows.append([r['dataset'].replace('_','-'),r['comparison'].replace('_',' '),int(r['n']),f"{r['mean_diff']:.3f}",f"[{r['ci95_low']:.3f}, {r['ci95_high']:.3f}]",f"{r['wilcoxon_p']:.4f}",f"{r['cohens_dz']:.2f}"])
add_table(doc,['Dataset','Comparison','n','Mean diff','95% CI','Wilcoxon p','Cohen dz'],rows,'Table 8. Paired statistical tests for ridge-readout frozen-feature accuracy. Source: results/frozen_readout_statistical_tests.csv.')
add_fig(doc, FIG/'mechanism_map_v2.png','Figure 15. SparseMechanismBench mechanism map. Source: figures/mechanism_map_v2.png.',width=5.8)

# Discussion
doc.add_heading('7. Discussion',1)
doc.add_paragraph("The new experiments strengthen the project from a benchmark comparison into a mechanism-isolation study. The key result is no longer merely that MLPs and CNNs beat simple Hebbian/Oja features. The stronger result is that sparse local plasticity separates sparsity from task discrimination. Oja features can preserve task-accessible information on MNIST/Fashion-MNIST when a ridge readout is used, but Sparse Oja and Homeostatic Sparse Oja maintain high sparsity while losing a large amount of class-separable structure. This means the failure is not purely a weak-readout artifact: sparse local features are geometrically less class-aligned.")
doc.add_paragraph("Homeostasis and competition improve the mechanism story but do not solve the problem. Homeostatic Sparse Oja sometimes improves over fixed Sparse Oja, particularly in full benchmark digits and Fashion-MNIST conditions, but it does not close the gap to supervised models. The sparsity sweep shows why: extreme sparsity can damage task discrimination. Replay helps retention in continual-learning subsets, but replay alone does not produce task-discriminative representations. These results support a hybrid interpretation: biological learning efficiency likely emerges from coordinated mechanisms, including local plasticity, sparse coding, competition, homeostasis, replay, feedback, and task-relevant signals.")
doc.add_paragraph("The negative result matters scientifically. It challenges a common oversimplified explanation of brain-like efficiency: sparsity by itself. Sparse local plasticity is useful, but the experiments show that it is insufficient when isolated from stronger credit-assignment or task-relevant feedback mechanisms.")

# Limitations
doc.add_heading('8. Limitations',1)
for item in [
"Homeostatic Sparse Oja uses a simple adaptive threshold rule, not a biologically complete model of neuromodulation, dendritic computation, or recurrent circuitry.",
"Frozen-readout and representation-geometry experiments use fixed mechanism-isolation subsets; they are not state-of-the-art full-data benchmarks.",
"MLP and CNN hidden-feature readouts were run with five seeds, while local/raw feature readouts were run with ten seeds. This is reported explicitly in the tables.",
"Approximate MACs are computational proxies and should not be interpreted as measured energy consumption.",
"Fashion-MNIST and MNIST results use the uploaded IDX files, but future work should publish a stable GitHub/OSF/Zenodo archive with versioned data-loading instructions.",
"Reward-modulated Hebbian learning, feedback alignment hybrids, recurrent predictive coding, and EWC were not fully implemented in this pass."
]: doc.add_paragraph(item, style='List Bullet')

# Conclusion
doc.add_heading('9. Conclusion',1)
doc.add_paragraph("SparseMechanismBench tests whether sparse local plasticity is sufficient for task-discriminative learning. The answer, under these benchmarks, is no. Sparse Oja-style learning reliably produces sparse representations, but sparsity alone does not guarantee class separability, strong readout performance, or stable continual learning. Oja features can preserve useful information, especially with ridge readouts, but imposing extreme sparsity sharply reduces task-accessible structure. Homeostasis and replay partially change the trade-off but do not replace task-directed credit assignment. The strongest conclusion is that biological learning efficiency likely depends on interacting mechanisms rather than a single local update rule: local plasticity, sparse coding, competition, homeostasis, replay, feedback, and task-relevant signals must work together.")

# Code availability
doc.add_heading('10. Code and Data Availability',1)
doc.add_paragraph("The accompanying package contains a reproducible codebase with README.md, requirements.txt, data_loading.py, models.py, train_classification.py, train_hebbian_oja.py, train_sparse_oja.py, train_homeostatic_oja.py, run_sample_efficiency.py, run_continual_learning.py, run_full_benchmark.py, run_frozen_readout.py, run_representation_geometry.py, run_frozen_readout_geometry_uploaded.py, make_enhanced_figures.py, make_all_figures.py, raw CSV files, summary CSV files, generated PNG figures, and manuscript-generation scripts. Commands are documented in README.md. A public GitHub/OSF/Zenodo link should be added before preprint submission. No test-set tuning was performed. All reported numerical results must be traceable to raw CSV files and plotting scripts.")

# Project summary
doc.add_heading('11. One-Page STS Project Summary',1)
doc.add_paragraph("Problem: Modern AI learns through global gradients, while brains learn largely through local plasticity and sparse activity. Hypothesis: sparse local plasticity may improve efficiency and stability, but sparsity alone may not create task-discriminative representations. Methods: SparseMechanismBench compares supervised baselines, Oja, Sparse Oja, Homeostatic Sparse Oja, Standard CNN, sparsity sweeps, replay sweeps, frozen readouts, geometry metrics, and statistical tests across digits, MNIST, and Fashion-MNIST. Results: supervised models dominate classification; Oja features contain task-accessible information; Sparse Oja creates high sparsity but weaker class geometry; homeostasis/replay help partially but do not close the gap. Significance: the project isolates why sparsity alone is not enough and motivates hybrid biologically inspired learning.")

# Descriptions and originality
doc.add_heading('12. STS / College Application Descriptions',1)
doc.add_paragraph("150 characters: Built SparseMechanismBench showing sparse Oja learning creates sparse codes but needs readout/replay/homeostasis for task learning.")
doc.add_paragraph("350 characters: Designed a reproducible multi-dataset mechanism-isolation study comparing backpropagation, Oja, Sparse Oja, Homeostatic Sparse Oja, CNNs, replay, and frozen readouts. Results show sparsity alone is not enough: local rules create sparse representations but need additional mechanisms for task discrimination and stability.")
doc.add_heading('Originality Statement',2)
doc.add_paragraph("This project is original because it does not simply compare Hebbian/Oja learning against backpropagation. It isolates mechanisms: sparsity, competition, homeostasis, replay, readout strength, and representation geometry. The negative result is scientifically useful because it shows that representational sparsity and task-discriminative learning are separable properties.")
doc.add_heading('Reviewer-Response Memo',2)
doc.add_paragraph("The project was improved by adding a full reproducibility codebase, replacing the weak compact CNN concern with a LeNet-style Standard CNN baseline, adding Homeostatic Sparse Oja, running sparsity and replay sweeps, adding frozen-readout experiments on uploaded MNIST/Fashion-MNIST IDX files, adding representation-geometry and feature-health diagnostics, adding statistical tests, and rewriting the paper around the mechanism-level claim that sparsity alone is insufficient.")

# References
doc.add_heading('References',1)
refs=[
"Bengio, Y. (2014). How auto-encoders could provide credit assignment in deep networks via target propagation. arXiv:1407.7906.",
"Bi, G.-Q., & Poo, M.-M. (1998). Synaptic modifications in cultured hippocampal neurons: Dependence on spike timing, synaptic strength, and postsynaptic cell type. Journal of Neuroscience, 18(24), 10464-10472. https://doi.org/10.1523/JNEUROSCI.18-24-10464.1998",
"Bienenstock, E. L., Cooper, L. N., & Munro, P. W. (1982). Theory for the development of neuron selectivity: Orientation specificity and binocular interaction in visual cortex. Journal of Neuroscience, 2(1), 32-48. https://doi.org/10.1523/JNEUROSCI.02-01-00032.1982",
"Hebb, D. O. (1949). The organization of behavior: A neuropsychological theory. Wiley.",
"Kirkpatrick, J., Pascanu, R., Rabinowitz, N., et al. (2017). Overcoming catastrophic forgetting in neural networks. Proceedings of the National Academy of Sciences, 114(13), 3521-3526. https://doi.org/10.1073/pnas.1611835114",
"LeCun, Y., Bottou, L., Bengio, Y., & Haffner, P. (1998). Gradient-based learning applied to document recognition. Proceedings of the IEEE, 86(11), 2278-2324. https://doi.org/10.1109/5.726791",
"Lillicrap, T. P., Cownden, D., Tweed, D. B., & Akerman, C. J. (2016). Random synaptic feedback weights support error backpropagation for deep learning. Nature Communications, 7, 13276. https://doi.org/10.1038/ncomms13276",
"McClelland, J. L., McNaughton, B. L., & O'Reilly, R. C. (1995). Why there are complementary learning systems in the hippocampus and neocortex. Psychological Review, 102(3), 419-457. https://doi.org/10.1037/0033-295X.102.3.419",
"McCloskey, M., & Cohen, N. J. (1989). Catastrophic interference in connectionist networks: The sequential learning problem. In G. H. Bower (Ed.), The psychology of learning and motivation (Vol. 24, pp. 109-165). Academic Press.",
"Oja, E. (1982). A simplified neuron model as a principal component analyzer. Journal of Mathematical Biology, 15(3), 267-273. https://doi.org/10.1007/BF00275687",
"Olshausen, B. A., & Field, D. J. (1996). Emergence of simple-cell receptive field properties by learning a sparse code for natural images. Nature, 381, 607-609. https://doi.org/10.1038/381607a0",
"Rao, R. P. N., & Ballard, D. H. (1999). Predictive coding in the visual cortex: A functional interpretation of some extra-classical receptive-field effects. Nature Neuroscience, 2, 79-87. https://doi.org/10.1038/4580",
"Rumelhart, D. E., Hinton, G. E., & Williams, R. J. (1986). Learning representations by back-propagating errors. Nature, 323, 533-536. https://doi.org/10.1038/323533a0",
"Scellier, B., & Bengio, Y. (2017). Equilibrium propagation: Bridging the gap between energy-based models and backpropagation. Frontiers in Computational Neuroscience, 11, 24. https://doi.org/10.3389/fncom.2017.00024",
"Whittington, J. C. R., & Bogacz, R. (2019). Theories of error back-propagation in the brain. Trends in Cognitive Sciences, 23(3), 235-250. https://doi.org/10.1016/j.tics.2018.12.005",
"Xiao, H., Rasul, K., & Vollgraf, R. (2017). Fashion-MNIST: A novel image dataset for benchmarking machine learning algorithms. arXiv:1708.07747."
]
for ref in refs: doc.add_paragraph(ref)

# Save
outdoc=OUT/'SparseMechanismBench_final_manuscript.docx'
doc.save(outdoc)
print(outdoc)
